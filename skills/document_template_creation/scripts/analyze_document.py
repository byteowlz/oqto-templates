#!/usr/bin/env python3
"""
Analyze existing documents (docx, pdf) and extract structure for tmpltr templates.

Usage:
    python analyze_document.py input.docx --output analysis.json
    python analyze_document.py input.pdf --format toml --output content.toml
"""

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any


def analyze_docx(file_path: Path) -> dict:
    """Analyze a DOCX file and extract structure."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    doc = Document(file_path)
    analysis = {
        "file_type": "docx",
        "file_name": file_path.name,
        "paragraphs": [],
        "tables": [],
        "placeholders": [],
        "headings": [],
    }

    # Extract paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            # Detect placeholders like [NAME], {date}, ___, etc.
            placeholders = re.findall(r"\[([^\]]+)\]|\{([^}]+)\}|_{3,}|<([^>]+)>", text)
            for match in placeholders:
                placeholder = next(p for p in match if p)
                analysis["placeholders"].append({
                    "text": f"[{placeholder}]" if match[0] else (f"{{{placeholder}}}" if match[1] else placeholder),
                    "context": text[:100],
                    "paragraph_index": i,
                })

            # Detect headings
            if para.style.name.startswith("Heading"):
                level = int(para.style.name.replace("Heading ", "")) if para.style.name.replace("Heading ", "").isdigit() else 1
                analysis["headings"].append({
                    "level": level,
                    "text": text,
                    "index": i,
                })
            else:
                analysis["paragraphs"].append({
                    "index": i,
                    "text": text[:200],
                    "style": para.style.name,
                })

    # Extract tables
    for i, table in enumerate(doc.tables):
        table_data = {
            "index": i,
            "rows": len(table.rows),
            "columns": len(table.columns),
            "header": [],
            "sample_data": [],
        }

        if table.rows:
            # First row as header
            table_data["header"] = [cell.text.strip() for cell in table.rows[0].cells]

            # Sample data (first 3 rows)
            for row in table.rows[1:4]:
                table_data["sample_data"].append([cell.text.strip() for cell in row.cells])

        analysis["tables"].append(table_data)

    return analysis


def analyze_pdf(file_path: Path) -> dict:
    """Analyze a PDF file and extract structure."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return {"error": "PyMuPDF not installed. Run: pip install PyMuPDF"}

    doc = fitz.open(file_path)
    analysis = {
        "file_type": "pdf",
        "file_name": file_path.name,
        "pages": len(doc),
        "text_blocks": [],
        "placeholders": [],
        "tables": [],
    }

    for page_num in range(min(3, len(doc))):  # Analyze first 3 pages
        page = doc[page_num]
        text = page.get_text()

        # Detect placeholders
        placeholders = re.findall(r"\[([^\]]+)\]|\{([^}]+)\}|_{3,}|<([^>]+)>", text)
        for match in placeholders:
            placeholder = next(p for p in match if p)
            analysis["placeholders"].append({
                "text": f"[{placeholder}]",
                "page": page_num + 1,
            })

        # Text blocks (first few lines per page)
        lines = [l.strip() for l in text.split("\n") if l.strip()][:20]
        analysis["text_blocks"].append({
            "page": page_num + 1,
            "lines": lines,
        })

    doc.close()
    return analysis


def generate_toml(analysis: dict) -> str:
    """Generate a TOML content file structure from analysis."""
    toml_lines = ["# Generated content structure", "# Review and customize as needed", ""]

    # Meta section
    toml_lines.extend([
        "[meta]",
        'template = "template.typ"',
        f'template_id = "{analysis.get("file_name", "document").replace(".", "_")}"',
        "template_version = \"1.0.0\"",
        "",
    ])

    # Document section based on detected structure
    toml_lines.extend([
        "[document]",
        f'title = ""  # Detected from document',
        f'date = ""  # e.g., 2025-01-15',
        "",
    ])

    # Headings become potential sections
    for heading in analysis.get("headings", [])[:5]:
        key = re.sub(r"[^a-zA-Z0-9]", "_", heading["text"].lower())[:30]
        toml_lines.extend([
            f"[blocks.{key}]",
            f'title = "{heading["text"]}"',
            'format = "markdown"',
            'content = """',
            f"Content for {heading['text']} section",
            '"""',
            "",
        ])

    # Placeholders become fields
    if analysis.get("placeholders"):
        toml_lines.extend([
            "# Detected placeholders - map to appropriate sections",
            "",
        ])
        seen = set()
        for ph in analysis["placeholders"]:
            key = ph["text"].strip("[]{}<>").lower().replace(" ", "_")
            if key not in seen:
                seen.add(key)
                toml_lines.append(f"# {key} = \"{ph['text']}\"  # {ph.get('context', '')[:50]}")

    # Tables
    if analysis.get("tables"):
        toml_lines.extend([
            "",
            "# Detected tables",
        ])
        for i, table in enumerate(analysis["tables"]):
            toml_lines.extend([
                f"",
                f"# Table {i+1}: {table.get('columns', '?')} columns",
            ])

    return "\n".join(toml_lines)


def generate_typst(analysis: dict) -> str:
    """Generate a starter Typst template from analysis."""
    lines = [
        "// Template generated from document analysis",
        "// @description: Template based on " + analysis.get("file_name", "source document"),
        "// @version: 1.0.0",
        "",
        '#import "@local/tmpltr-lib:1.0.0": tmpltr-data, get',
        "",
        "#let data = tmpltr-data()",
        "",
        "// Page setup",
        '#set page(paper: "a4", margin: 2.5cm)',
        '#set text(font: "Helvetica Neue", size: 11pt)',
        "",
    ]

    # Add fields based on placeholders
    placeholders = analysis.get("placeholders", [])
    if placeholders:
        lines.append("// Document fields")
        seen = set()
        for ph in placeholders[:10]:  # First 10 placeholders
            key = ph["text"].strip("[]{}<>").lower().replace(" ", "_")
            if key not in seen:
                seen.add(key)
                lines.append(f'#let {key} = get(data, "document.{key}", default: "{ph["text"]}")')
        lines.append("")

    # Add headings as section placeholders
    for heading in analysis.get("headings", [])[:5]:
        key = re.sub(r"[^a-zA-Z0-9]", "_", heading["text"].lower())[:30]
        lines.append(f'// Section: {heading["text"]}')
        lines.append(f'#get(data, "blocks.{key}.content", default: "")')
        lines.append("")

    # Add table placeholder
    if analysis.get("tables"):
        lines.append("// Tables (requires customization)")
        lines.append("#let table_data = get(data, \"table\", default: ())")
        lines.append("")

    lines.append("// End of template")
    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(
        description="Analyze documents and extract structure for tmpltr"
    )
    parser.add_argument("input", type=Path, help="Input document (docx or pdf)")
    parser.add_argument("--output", "-o", type=Path, help="Output file")
    parser.add_argument(
        "--format",
        "-f",
        choices=["json", "toml", "typst"],
        default="json",
        help="Output format",
    )

    args = parser.parse_args()

    if not args.input.exists():
        print(f"Error: File not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    # Determine file type
    suffix = args.input.suffix.lower()

    if suffix == ".docx":
        analysis = analyze_docx(args.input)
    elif suffix == ".pdf":
        analysis = analyze_pdf(args.input)
    else:
        print(f"Error: Unsupported file type: {suffix}", file=sys.stderr)
        sys.exit(1)

    # Generate output
    if args.format == "json":
        output = json.dumps(analysis, indent=2, ensure_ascii=False)
    elif args.format == "toml":
        output = generate_toml(analysis)
    elif args.format == "typst":
        output = generate_typst(analysis)
    else:
        output = json.dumps(analysis, indent=2)

    # Write or print output
    if args.output:
        args.output.write_text(output, encoding="utf-8")
        print(f"Wrote {args.format} output to {args.output}")
    else:
        print(output)


if __name__ == "__main__":
    main()
